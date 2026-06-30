"""Email attachment document-import helpers."""

from __future__ import annotations

import os
import shutil
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any

from fastapi import Request


def attachment_basename(filepath: Path) -> tuple[str, str, str] | dict[str, str]:
    base = Path(filepath).name
    if base.startswith("."):
        return {"error": "Invalid filename", "filename": base}
    ext = Path(base).suffix.lower()
    title = os.path.splitext(base)[0]
    return base, ext, title


def tag_document_with_source(
    doc_id: str,
    *,
    uid: str,
    folder: str,
    account_id: str | None,
    message_id: str,
    logger: Any = None,
) -> None:
    if not doc_id:
        return
    try:
        from src.database import Document as DbDocument
        from src.database import SessionLocal

        db = SessionLocal()
        try:
            doc = db.query(DbDocument).filter(DbDocument.id == doc_id).first()
            if doc:
                doc.source_email_uid = str(uid)
                doc.source_email_folder = folder
                doc.source_email_account_id = account_id or ""
                doc.source_email_message_id = message_id
                db.commit()
        finally:
            db.close()
    except Exception as exc:
        if logger is not None:
            logger.warning(f"tag doc source-email failed: {exc}")


def resolve_document_session(request: Request, logger: Any = None) -> str | None:
    try:
        from src.auth_helpers import get_current_user
        from src.database import Session as DbSession
        from src.database import SessionLocal

        doc_user = get_current_user(request)
        db = SessionLocal()
        try:
            query = db.query(DbSession)
            if doc_user:
                query = query.filter(DbSession.owner == doc_user)
            session = query.order_by(DbSession.updated_at.desc()).first()
            return session.id if session else None
        finally:
            db.close()
    except Exception as exc:
        if logger is not None:
            logger.warning(f"resolve doc session failed: {exc}")
        return None


def docx_to_markdown(filepath: Path) -> str | dict[str, str]:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        return {"error": "python-docx not installed", "filename": filepath.name}
    try:
        docx = DocxDocument(str(filepath))
    except Exception as exc:
        return {"error": f"Failed to read docx: {exc}", "filename": filepath.name}

    lines: list[str] = []
    for paragraph in docx.paragraphs:
        text = paragraph.text or ""
        style = (paragraph.style.name if paragraph.style else "") or ""
        if not text.strip():
            lines.append("")
            continue
        if style.startswith("Heading 1"):
            lines.append(f"# {text}")
        elif style.startswith("Heading 2"):
            lines.append(f"## {text}")
        elif style.startswith("Heading 3"):
            lines.append(f"### {text}")
        elif style.startswith("Heading "):
            lines.append(f"#### {text}")
        elif style.startswith("List Bullet"):
            lines.append(f"- {text}")
        elif style.startswith("List Number"):
            lines.append(f"1. {text}")
        else:
            lines.append(text)

    for table in docx.tables:
        lines.append("")
        for row_index, row in enumerate(table.rows):
            cells = [
                (cell.text or "").replace("|", "\\|").replace("\n", " ").strip()
                for cell in row.cells
            ]
            lines.append("| " + " | ".join(cells) + " |")
            if row_index == 0:
                lines.append("|" + "|".join(["---"] * len(cells)) + "|")
        lines.append("")

    return "\n".join(lines).strip() or f"_(empty {filepath.name})_"


def create_markdown_attachment_document(
    *,
    title: str,
    content: str,
    session_id: str | None,
    summary: str,
) -> str:
    from src.database import Document as DbDocument
    from src.database import DocumentVersion
    from src.database import SessionLocal

    doc_id = str(uuid.uuid4())
    version_id = str(uuid.uuid4())
    db = SessionLocal()
    try:
        db.query(DbDocument).filter(DbDocument.is_active == True).update({"is_active": False})
        db.add(
            DbDocument(
                id=doc_id,
                session_id=session_id,
                title=title,
                language="markdown",
                current_content=content,
                version_count=1,
                is_active=True,
            )
        )
        db.add(
            DocumentVersion(
                id=version_id,
                document_id=doc_id,
                version_number=1,
                content=content,
                summary=summary,
                source="upload",
            )
        )
        db.commit()
    finally:
        db.close()
    return doc_id


def create_pdf_attachment_document(filepath: Path, *, title: str, session_id: str | None, logger: Any = None) -> str | None:
    from src.constants import UPLOAD_DIR
    from src.pdf_form_doc import (
        create_form_markdown_document,
        create_plain_pdf_document,
        save_field_sidecar,
    )
    from src.pdf_forms import extract_fields, has_form_fields

    upload_id = f"{uuid.uuid4().hex}.pdf"
    today = datetime.utcnow().strftime("%Y/%m/%d")
    dated_dir = os.path.join(UPLOAD_DIR, today)
    os.makedirs(dated_dir, exist_ok=True)
    dest_path = os.path.join(dated_dir, upload_id)
    shutil.copyfile(str(filepath), dest_path)

    is_form = False
    try:
        is_form = has_form_fields(dest_path)
    except Exception as exc:
        if logger is not None:
            logger.warning(f"has_form_fields failed for attachment PDF: {exc}")

    if is_form:
        fields = extract_fields(dest_path)
        save_field_sidecar(dest_path, fields)
        return create_form_markdown_document(
            session_id=session_id,
            fields=fields,
            upload_id=upload_id,
            title=title,
            intro_text=None,
        )

    return create_plain_pdf_document(
        session_id=session_id,
        upload_id=upload_id,
        title=title,
    )


def attachment_as_document_response(
    filepath: Path,
    msg,
    *,
    uid: str,
    folder: str,
    account_id: str | None,
    request: Request,
    logger: Any = None,
) -> dict[str, Any]:
    parsed = attachment_basename(filepath)
    if isinstance(parsed, dict):
        return parsed
    base, ext, title = parsed
    source_message_id = (msg.get("Message-ID") or "").strip()

    if ext not in (".pdf", ".docx", ".txt", ".md", ".markdown"):
        return {"error": f"Unsupported attachment type: {ext}", "filename": base}

    doc_session_id = resolve_document_session(request, logger=logger)

    if ext == ".pdf":
        doc_id = create_pdf_attachment_document(filepath, title=title, session_id=doc_session_id, logger=logger)
        if not doc_id:
            return {"error": "Failed to create document"}
        tag_document_with_source(
            doc_id,
            uid=uid,
            folder=folder,
            account_id=account_id,
            message_id=source_message_id,
            logger=logger,
        )
        return {"doc_id": doc_id, "filename": filepath.name}

    if ext == ".docx":
        content = docx_to_markdown(filepath)
        if isinstance(content, dict):
            return content
        doc_id = create_markdown_attachment_document(
            title=title,
            content=content,
            session_id=doc_session_id,
            summary="Imported from DOCX",
        )
    else:
        try:
            content = filepath.read_text(encoding="utf-8", errors="replace")
        except Exception as exc:
            return {"error": f"Failed to read text file: {exc}", "filename": base}
        doc_id = create_markdown_attachment_document(
            title=title,
            content=content,
            session_id=doc_session_id,
            summary="Imported from email attachment",
        )

    tag_document_with_source(
        doc_id,
        uid=uid,
        folder=folder,
        account_id=account_id,
        message_id=source_message_id,
        logger=logger,
    )
    return {"doc_id": doc_id, "filename": filepath.name}
